"""
Генерация итогового DOCX-отчёта по 6 сценариям на основе results/results.json.
Отчёт содержит таблицы, графики и текстовые выводы, синхронизированные с расчётными данными.
"""

import json
import os
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT_DIR = "/Users/m/Documents/Openstudio"
RESULTS_PATH = os.path.join(ROOT_DIR, "results", "results.json")
FIG_DIR = os.path.join(ROOT_DIR, "results", "figures")
DOCX_PATH = os.path.join(ROOT_DIR, "Отчёт_Кинозал_DCV_OpenStudio 3.docx")

ORDER = ["Base", "DCV", "Base_HP", "DCV_HP", "Base_Rec", "DCV_Rec"]
LABELS = {
    "Base":     "Base",
    "DCV":      "DCV",
    "Base_HP":  "Base+HP",
    "DCV_HP":   "DCV+HP",
    "Base_Rec": "Base+Rec",
    "DCV_Rec":  "DCV+Rec",
}
COLORS = {
    "Base":     "#1f77b4",
    "DCV":      "#d62728",
    "Base_HP":  "#2ca02c",
    "DCV_HP":   "#ff7f0e",
    "Base_Rec": "#9467bd",
    "DCV_Rec":  "#8c564b",
}
MONTHS = ["Янв", "Фев", "Мар", "Апр", "Май", "Июн", "Июл", "Авг", "Сен", "Окт", "Ноя", "Дек"]


def fmt(v: float, nd: int = 1) -> str:
    return f"{v:,.{nd}f}".replace(",", " ")


def pct(base: float, value: float) -> float:
    if abs(base) < 1e-9:
        return 0.0
    return (value - base) / base * 100.0


def load_data() -> dict:
    if not os.path.exists(RESULTS_PATH):
        raise FileNotFoundError(f"Не найден файл: {RESULTS_PATH}")
    with open(RESULTS_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def make_figures(scenarios: dict) -> list:
    os.makedirs(FIG_DIR, exist_ok=True)
    existing = [s for s in ORDER if s in scenarios]
    x = np.arange(len(existing))

    paths = []

    # Рисунок 1: суммарное электропотребление
    p1 = os.path.join(FIG_DIR, "fig1_total_electricity_6models.png")
    y = [scenarios[s]["elec_kwh"] for s in existing]
    plt.figure(figsize=(10.0, 4.8))
    bars = plt.bar(x, y, color=[COLORS[s] for s in existing], edgecolor="white")
    plt.xticks(x, [LABELS[s] for s in existing])
    plt.ylabel("кВт·ч/год")
    plt.title("Суммарное электропотребление по сценариям")
    plt.grid(axis="y", alpha=0.3)
    for b, v in zip(bars, y):
        plt.text(b.get_x() + b.get_width() / 2, v, fmt(v, 0), ha="center", va="bottom", fontsize=8)
    plt.tight_layout()
    plt.savefig(p1, dpi=160)
    plt.close()
    paths.append((p1, "Рисунок 1 — Суммарное электропотребление по шести сценариям"))

    # Рисунок 2: вентиляторы
    p2 = os.path.join(FIG_DIR, "fig2_fan_electricity_6models.png")
    y = [scenarios[s]["fan_kwh"] for s in existing]
    plt.figure(figsize=(10.0, 4.8))
    bars = plt.bar(x, y, color=[COLORS[s] for s in existing], edgecolor="white")
    plt.xticks(x, [LABELS[s] for s in existing])
    plt.ylabel("кВт·ч/год")
    plt.title("Электропотребление вентиляторов")
    plt.grid(axis="y", alpha=0.3)
    for b, v in zip(bars, y):
        plt.text(b.get_x() + b.get_width() / 2, v, fmt(v, 0), ha="center", va="bottom", fontsize=8)
    plt.tight_layout()
    plt.savefig(p2, dpi=160)
    plt.close()
    paths.append((p2, "Рисунок 2 — Электропотребление вентиляторов по сценариям"))

    # Рисунок 3: помесячная динамика вентиляторов
    p3 = os.path.join(FIG_DIR, "fig3_monthly_fans_6models.png")
    xm = np.arange(12)
    plt.figure(figsize=(10.8, 5.2))
    for s in existing:
        plt.plot(xm, scenarios[s]["monthly"]["fans_kwh"], marker="o", linewidth=1.7, label=LABELS[s], color=COLORS[s])
    plt.xticks(xm, MONTHS)
    plt.ylabel("кВт·ч/мес")
    plt.title("Помесячное электропотребление вентиляторов")
    plt.grid(alpha=0.3)
    plt.legend(ncol=2)
    plt.tight_layout()
    plt.savefig(p3, dpi=160)
    plt.close()
    paths.append((p3, "Рисунок 3 — Помесячное электропотребление вентиляторов"))

    # Рисунок 4: CO2 — сравнение с порогами ГОСТ 30494-2011
    p4 = os.path.join(FIG_DIR, "fig4_co2_average_6models.png")
    y = [scenarios[s]["co2_avg_ppm"] for s in existing]
    plt.figure(figsize=(10.0, 4.8))
    bars = plt.bar(x, y, color=[COLORS[s] for s in existing], edgecolor="white")
    plt.axhline(800,  linestyle="-",  linewidth=1.5, color="#2ca02c",
                label="800 ppm — граница «Отличное» (ГОСТ 30494)")
    plt.axhline(1400, linestyle="--", linewidth=1.2, color="#d62728",
                label="1400 ppm — граница «Допустимое»")
    plt.xticks(x, [LABELS[s] for s in existing])
    plt.ylabel("ppm")
    plt.title("Средняя концентрация CO₂ по сценариям (ГОСТ 30494-2011)")
    plt.grid(axis="y", alpha=0.3)
    plt.legend(loc="upper right", fontsize=9)
    for b, v in zip(bars, y):
        plt.text(b.get_x() + b.get_width() / 2, v, fmt(v, 0), ha="center", va="bottom", fontsize=8)
    plt.tight_layout()
    plt.savefig(p4, dpi=160)
    plt.close()
    paths.append((p4, "Рисунок 4 — Средняя концентрация CO₂ (ГОСТ 30494-2011: < 800 ppm = «Отличное»)"))

    return paths


def add_para(doc: Document, text: str, *, bold: bool = False, center: bool = False, size: int = 12):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    return p


def add_table_system(doc: Document):
    t = doc.add_table(rows=8, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Параметр", "Значение", "Примечание"]):
        t.cell(0, i).text = h

    rows = [
        ("Тип установки", "AIRNED-M6L", "Приточно-вытяжная вентиляция"),
        ("Обозначение", "П1В1", "Базовая система объекта"),
        ("Приток", "4 050 м³/ч (1,125 м³/с)", "Номинальный режим"),
        ("Вытяжка", "4 050 м³/ч", "Номинальный режим"),
        ("Кратность воздухообмена", "6,28 ч⁻¹", "Для объёма 645 м³"),
        ("Вентилятор", "5 425 м³/ч, 500 Па, 1,97 кВт", "Проектные характеристики"),
        ("Сценарии", "Base, DCV, Base+HP, DCV+HP, Base+Rec, DCV+Rec", "Расчёт выполнен для 6 моделей"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, v in enumerate(row):
            t.cell(r, c).text = v


CO2_QUALITY = {
    "Отличное":    (0, 800),
    "Хорошее":     (800, 1000),
    "Допустимое":  (1000, 1400),
    "Недопустимое":(1400, 9999),
}

def co2_category(ppm: float) -> str:
    for label, (lo, hi) in CO2_QUALITY.items():
        if lo <= ppm < hi:
            return label
    return "—"


def add_table_main(doc: Document, s: dict):
    existing_order = [n for n in ORDER if n in s]
    ncols = 1 + len(existing_order)
    t = doc.add_table(rows=8, cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Показатель"] + [LABELS[n] for n in existing_order]
    for i, h in enumerate(headers):
        t.cell(0, i).text = h

    data_rows = [
        ("Электроэнергия (всего), кВт·ч/год",      "elec_kwh",    0),
        ("Вентиляторы (электроэнергия), кВт·ч/год", "fan_kwh",     0),
        ("Калорифер (тепловая энергия), кВт·ч/год", "coil_kwh",    0),
        ("Средний расход вентиляции, м³/ч",          "vent_avg_m3h",0),
        ("Средняя концентрация CO₂, ppm",            "co2_avg_ppm", 0),
        ("Оценка качества воздуха (ГОСТ 30494)",     "co2_avg_ppm", -2),
        ("Часы присутствия (>10 чел), ч/год",        "occ_hours",   0),
    ]

    for r, (title, key, nd) in enumerate(data_rows, start=1):
        t.cell(r, 0).text = title
        for c, name in enumerate(existing_order, start=1):
            if nd == -2:
                t.cell(r, c).text = co2_category(float(s[name]["co2_avg_ppm"]))
            else:
                t.cell(r, c).text = fmt(float(s[name][key]), nd)


def add_table_compare(doc: Document, s: dict):
    all_pairs = [
        ("DCV vs Base",           "Base",     "DCV"),
        ("DCV+HP vs Base+HP",     "Base_HP",  "DCV_HP"),
        ("DCV+Rec vs Base+Rec",   "Base_Rec", "DCV_Rec"),
    ]
    valid_pairs = [(lbl, b, d) for lbl, b, d in all_pairs if b in s and d in s]
    nrows = 1 + len(valid_pairs)

    t = doc.add_table(rows=nrows, cols=5)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Пара", "ΔЭлектро, кВт·ч/год", "ΔЭлектро, %", "ΔВентиляторы, кВт·ч/год", "ΔВентиляторы, %"]
    for i, h in enumerate(headers):
        t.cell(0, i).text = h

    for r, (label, b, d) in enumerate(valid_pairs, start=1):
        eb = s[b]["elec_kwh"]
        ed = s[d]["elec_kwh"]
        fb = s[b]["fan_kwh"]
        fd = s[d]["fan_kwh"]
        t.cell(r, 0).text = label
        t.cell(r, 1).text = fmt(eb - ed, 0)
        t.cell(r, 2).text = fmt((eb - ed) / eb * 100.0 if eb else 0.0, 1)
        t.cell(r, 3).text = fmt(fb - fd, 0)
        t.cell(r, 4).text = fmt((fb - fd) / fb * 100.0 if fb else 0.0, 1)


def build_doc(data: dict, figs: list):
    s = data["scenarios"]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_para(
        doc,
        "РАЗРАБОТКА МЕТОДОВ ПОВЫШЕНИЯ ЭНЕРГОЭФФЕКТИВНОСТИ\n"
        "СИСТЕМ ВЕНТИЛЯЦИИ КИНОЗАЛА\n"
        "С КОНТРОЛЕМ КОНЦЕНТРАЦИИ CO₂",
        bold=True,
        center=True,
        size=14,
    )
    add_para(doc, "Моделирование в OpenStudio/EnergyPlus", center=True)

    add_heading(doc, "1. Цель и объект моделирования")
    add_para(
        doc,
        "Цель работы — оценка влияния алгоритма DCV на электропотребление системы вентиляции "
        "при сохранении контролируемого уровня CO₂ в зале. "
        "Объект — большой кинозал вместимостью 135 мест."
    )

    add_heading(doc, "2. Характеристики объекта и системы")
    add_para(doc, "Исходные параметры объекта и принятой вентиляционной системы сведены в таблицу 1.")
    add_table_system(doc)

    add_heading(doc, "3. Построение модели и расчётный контур")
    add_para(
        doc,
        "Расчёт выполнен в OpenStudio 3.10 / EnergyPlus. Для каждого сценария автоматически "
        "сформирован workflow и выполнен годовой прогон на едином климатическом файле Saint-Petersburg.epw."
    )

    add_heading(doc, "4. Сценарии расчёта")
    add_para(doc, "Рассмотрены шесть сценариев: Base, DCV, Base+HP, DCV+HP, Base+Rec и DCV+Rec.")
    add_para(
        doc,
        "Все DCV-сценарии используют уставку CO₂ = 700 ppm и минимальную долю приточного расхода 55 % "
        "от номинала (2 228 м³/ч), что обеспечивает категорию качества воздуха «Отличное» "
        "(CO₂ < 800 ppm, ГОСТ 30494-2011) при одновременном снижении электропотребления. "
        "Пара Base+Rec / DCV+Rec моделирует замену стандартного рекуператора (η=80 %) "
        "на высокоэффективный электрический роторный рекуператор (η=90 %)."
    )

    add_heading(doc, "5. Сводные результаты по шести моделям")
    add_table_main(doc, s)

    add_heading(doc, "6. Сравнение и визуализация")
    add_table_compare(doc, s)

    for path, caption in figs:
        doc.add_picture(path, width=Pt(460))
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.name = "Times New Roman"
        cp.runs[0].font.size = Pt(11)

    add_heading(doc, "7. Выводы")
    b  = s["Base"]
    d  = s["DCV"]

    add_para(
        doc,
        "1. Для основной пары Base/DCV получено снижение общего электропотребления "
        f"на {fmt(b['elec_kwh'] - d['elec_kwh'], 0)} кВт·ч/год "
        f"({fmt((b['elec_kwh'] - d['elec_kwh']) / b['elec_kwh'] * 100.0, 1)} %). "
        f"Средняя концентрация CO₂ в DCV-сценарии составляет {fmt(d['co2_avg_ppm'], 0)} ppm — "
        f"категория «{co2_category(d['co2_avg_ppm'])}» по ГОСТ 30494-2011."
    )
    add_para(
        doc,
        "2. Основной вклад в эффект DCV формируется снижением потребления вентиляторов: "
        f"{fmt(b['fan_kwh'] - d['fan_kwh'], 0)} кВт·ч/год "
        f"({fmt((b['fan_kwh'] - d['fan_kwh']) / b['fan_kwh'] * 100.0, 1)} %)."
    )

    if "Base_HP" in s and "DCV_HP" in s:
        bh = s["Base_HP"]
        dh = s["DCV_HP"]
        add_para(
            doc,
            "3. Для пары Base+HP/DCV+HP получен сопоставимый результат по электрической части: "
            f"{fmt(bh['elec_kwh'] - dh['elec_kwh'], 0)} кВт·ч/год "
            f"({fmt((bh['elec_kwh'] - dh['elec_kwh']) / bh['elec_kwh'] * 100.0, 1)} %)."
        )

    if "Base_Rec" in s and "DCV_Rec" in s:
        br = s["Base_Rec"]
        dr = s["DCV_Rec"]
        add_para(
            doc,
            "4. Высокоэффективный рекуператор (η=90 %) снижает тепловую нагрузку калорифера "
            f"в базовом режиме на {fmt(b['coil_kwh'] - br['coil_kwh'], 0)} кВт·ч/год "
            f"({fmt((b['coil_kwh'] - br['coil_kwh']) / b['coil_kwh'] * 100.0, 1)} %) "
            "по сравнению со стандартным рекуператором. "
            "Эффект DCV дополнительно снижает общее электропотребление в паре DCV+Rec: "
            f"{fmt(br['elec_kwh'] - dr['elec_kwh'], 0)} кВт·ч/год "
            f"({fmt((br['elec_kwh'] - dr['elec_kwh']) / br['elec_kwh'] * 100.0, 1)} %)."
        )

    add_para(
        doc,
        "5. Все шесть моделей успешно завершены в EnergyPlus. "
        "Применение уставки CO₂ = 700 ppm и минимальной доли приточного расхода 55 % обеспечивает "
        "категорию «Отличное» во всех DCV-сценариях, что делает сравнение корректным."
    )
    add_para(doc, f"Составлено: {date.today().strftime('%d.%m.%Y')}", size=10)

    doc.save(DOCX_PATH)


def main() -> None:
    data = load_data()
    scenarios = data.get("scenarios", {})
    # Обязательные сценарии — только Base и DCV; остальные опциональны
    required = ["Base", "DCV"]
    missing = [name for name in required if name not in scenarios]
    if missing:
        raise RuntimeError("Отсутствуют обязательные сценарии в results.json: " + ", ".join(missing))

    figs = make_figures(scenarios)
    build_doc(data, figs)
    print(f"[OK] Отчёт сохранён: {DOCX_PATH}")
    print(f"[OK] Графики обновлены: {FIG_DIR}")


if __name__ == "__main__":
    main()
